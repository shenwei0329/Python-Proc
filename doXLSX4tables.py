#!/usr/local/bin/python2.7
# -*- coding: utf-8 -*-
#
#

from __future__ import unicode_literals

import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
import DataHandler.crWord
import DataHandler.doBox
from pylab import mpl

"""设置字符集
"""
reload(sys)
sys.setdefaultencoding('utf-8')

import xlrd
import time
from datetime import datetime


mpl.rcParams['font.sans-serif'] = ['SimHei']


Topic_lvl_number = 0
Topic = [u'一、',
         u'二、',
         u'三、',
         u'四、',
         u'五、',
         u'六、',
         u'七、',
         u'八、',
         u'九、',
         u'十、',
         u'十一、',
         u'十二、',
         ]

rec = []
task = {}
source = {}
plan = {}
t_type = {}
t_base = {}

doc = None


def _print(_str, title=False, title_lvl=0, color=None, align=None, paragrap=None ):

    global doc, Topic_lvl_number, Topic

    _str = u"%s" % _str.replace('\r', '').replace('\n','')

    _paragrap = None

    if title_lvl == 1:
        Topic_lvl_number = 0
    if title:
        if title_lvl==2:
            _str = Topic[Topic_lvl_number] + _str
            Topic_lvl_number += 1
        if align is not None:
            _paragrap = doc.addHead(_str, title_lvl, align=align)
        else:
            _paragrap = doc.addHead(_str, title_lvl)
    else:
        if align is not None:
            if paragrap is None:
                _paragrap = doc.addText(_str, color=color, align=align)
            else:
                doc.appendText(paragrap, _str, color=color, align=align)
        else:
            if paragrap is None:
                _paragrap = doc.addText(_str, color=color)
            else:
                doc.appendText(paragrap, _str, color=color)
    print(_str)

    return _paragrap


class XlsxHandler:

    def __init__(self, pathname):

        # logging.log(logging.WARN, u"XlsxHandler __init__(%s)" % pathname)

        self.data = xlrd.open_workbook(pathname)
        self.tables = self.data.sheets()
        self.table = self.tables[0]
        self.nrows = self.table.nrows

    def getSheetNumber(self):
        return len(self.tables)

    def setSheet(self, n):
        if n < len(self.tables):
            self.table = self.tables[n]
            self.nrows = self.table.nrows

    def getTableName(self):
        return self.table.name

    def getNrows(self):
        return self.nrows

    def getData(self, row, col):
        """
        获取xlsx记录单元的数据
        :param table: 数据源
        :param row: 行号
        :param col: 列号
        :return:
        """
        try:
            return self.table.row_values(row)[col]
        except:
            return None

    def getXlsxCol(self):

        # print(">>> getXlsxCol: %d" % self.table.ncols)
        return self.table.ncols

    def getXlsxColStr(self):

        print(">>> getXlsxColStr: %d" % self.table.ncols)
        _col = self.getXlsxColName(self.table.ncols)
        _str = ""
        for _s in _col:
            _str += (_s + ',')

        # _str = _str.decode("utf-8", "replace")
        return _str, self.table.ncols

    def getXlsxColName(self, nCol):

        _col = []
        for i in range(14, nCol):
            _colv = self.table.row_values(0)[i]
            _col.append(_colv)
            # print(">>> Col[%s]" % _colv)
        return _col

    def getXlsxAllColName(self, nCol):

        _col = []
        for i in range(nCol):
            _colv = self.table.row_values(0)[i]
            _col.append(_colv)
            # print(">>> Col[%s]" % _colv)
        return _col

    def getXlsxRow(self, i, nCol, lastRow):
        """
        获取某行数据
        :param i: 行号
        :param nCol: 字段数
        :param lastRow: 上一行数据
        :return: 指定行的数据
        """

        # print("%s- getXlsxRow[%d,%d]" % (time.ctime(), i, nCol))

        """ctype : 0 empty,1 string, 2 number, 3 date, 4 boolean, 5 error
        """
        _row = []
        for _i in range(nCol):
            __row = self.table.row_values(i)[_i]
            __ctype = self.table.cell(i, _i).ctype
            # print __row,__ctype
            if __ctype == 3:
                _date = datetime(*xlrd.xldate_as_tuple(__row, 0))
                __row = _date.strftime('%Y/%m/%d')
                __row = __row.split('/')
                __row = u"%d年%d月%d日" % (int(__row[0]), int(__row[1]), int(__row[2]))
            elif __ctype == 2:
                __row = str(__row)
                # __row = str(__row).replace('.0', '')
            elif __ctype == 5:
                __row = ''
            _row.append(__row)

        # print _row

        row = []
        _i = 0
        for _r in _row:

            # print(">>>[%s]" % _r)

            if _r is None or len(str(_r)) == 0:
                """用第一行的内容填充合并字段的其它单元"""
                if lastRow is not None:
                    _r = lastRow[_i]
            # print _r
            row.append(_r)
            _i = _i + 1
        return row


def doList(xlsx_handler):

    global rec

    _name = xlsx_handler.getTableName()
    _nrow = xlsx_handler.getNrows()

    # print "%s" % _name,_ncol

    for i in range(1, _nrow):
        _row = xlsx_handler.getXlsxRow(i, xlsx_handler.getXlsxCol(), None)
        # print _row
        _rec = {'member': _row[0], 'task': _row[1], 'source': _row[2], 'plan': _row[3], 'type': _row[4], 'base': _row[5]}
        rec.append(_rec)


def calInd(recs, out, title):

    for _r in recs:
        _v = _r[title].split(u'、')
        for _t in _v:
            if _t not in out:
                out[_t] = 0
            out[_t] += 1


def printTable(title, recs):

    _x = []
    _y = []
    _print(u'%s的统计如下：' % title)
    doc.addTable(1, 2, col_width=(4, 2))
    _title = (('text', u'数据标签'),
              ('text', u'统计值'))
    doc.addRow(_title)

    for _r in recs:
        if _r is "":
            continue
        _x.append(_r)
        _y.append(recs[_r])
        _text = (('text', _r), ('text', str(recs[_r])))
        doc.addRow(_text)

    doc.setTableFont(8)
    return _x, _y


def main(filename):

    global doc, rec, source, task, plan, t_type, t_base

    _cnt = 0

    """创建word文档实例
    """
    doc = DataHandler.crWord.createWord()
    """写入"主题"
    """
    doc.addHead(u'《个人工作情况》分析报告', 0, align=WD_ALIGN_PARAGRAPH.CENTER)

    _print('>>> 报告生成日期【%s】 <<<' % time.ctime(), align=WD_ALIGN_PARAGRAPH.CENTER)

    xlsx_handler = XlsxHandler(filename)

    try:

        _tables = xlsx_handler.getSheetNumber()

        for _idx in range(_tables):
            xlsx_handler.setSheet(_idx)
            _cnt += xlsx_handler.getNrows()
            doList(xlsx_handler)

        _print("I、总体情况", title=True, title_lvl=1)

        _print("数据来源", title=True, title_lvl=2)
        _print(u"参与的总人数：%d（人）" % len(rec))
        _print(u"参与的部门：")
        for _idx in range(_tables):
            xlsx_handler.setSheet(_idx)
            _dpt = xlsx_handler.getTableName()
            _dpt_member = xlsx_handler.getNrows() - 1
            _print(u"\t%s：%d（人）" % (_dpt, _dpt_member))

        _print("任务情况", title=True, title_lvl=2)
        calInd(rec, task, 'task')
        _x, _y = printTable("任务", task)
        _fn = DataHandler.doBox.doBarH(u"任务分布", "数量", _x, _y)
        doc.addPic(_fn, sizeof=5)

        _print("任务的来源情况", title=True, title_lvl=2)
        calInd(rec, source, 'source')
        _x, _y = printTable("任务来源", source)
        _fn = DataHandler.doBox.doBarH(u"任务来源分布", "数量", _x, _y)
        doc.addPic(_fn, sizeof=5)

        _print("任务的计划情况", title=True, title_lvl=2)
        calInd(rec, plan, 'plan')
        _x, _y = printTable("任务计划", plan)
        _fn = DataHandler.doBox.doBarH(u"任务计划分布", "数量", _x, _y)
        doc.addPic(_fn, sizeof=5)

        _print("任务的成果类型情况", title=True, title_lvl=2)
        calInd(rec, t_type, 'type')
        _x, _y = printTable("成果类型", t_type)
        _fn = DataHandler.doBox.doBarH(u"任务类型分布", "数量", _x, _y)
        doc.addPic(_fn, sizeof=5)

        _print("任务的成果归档情况", title=True, title_lvl=2)
        calInd(rec, t_base, 'base')
        _x, _y = printTable("成果归档", t_base)
        _fn = DataHandler.doBox.doBarH(u"任务归档分布", "数量", _x, _y)
        doc.addPic(_fn, sizeof=5)

        _print("II、部门情况", title=True, title_lvl=1)

        for _idx in range(_tables):
            xlsx_handler.setSheet(_idx)
            _dpt = xlsx_handler.getTableName()
            _dpt_member = xlsx_handler.getNrows() - 1

            _print(u"【%s】部门" % _dpt, title=True, title_lvl=2)

            rec = []
            task = {}
            source = {}
            plan = {}
            t_type = {}
            t_base = {}

            doList(xlsx_handler)

            _print("1、任务情况", title=True, title_lvl=3)
            calInd(rec, task, 'task')
            _x, _y = printTable("任务", task)
            _fn = DataHandler.doBox.doBarH(u"任务分布", "数量", _x, _y)
            doc.addPic(_fn, sizeof=5)

            _print("2、任务的来源情况", title=True, title_lvl=3)
            calInd(rec, source, 'source')
            _x, _y = printTable("任务来源", source)
            _fn = DataHandler.doBox.doBarH(u"任务来源分布", "数量", _x, _y)
            doc.addPic(_fn, sizeof=5)

            _print("3、任务的计划情况", title=True, title_lvl=3)
            calInd(rec, plan, 'plan')
            _x, _y = printTable("任务计划", plan)
            _fn = DataHandler.doBox.doBarH(u"任务计划分布", "数量", _x, _y)
            doc.addPic(_fn, sizeof=5)

            _print("4、任务的成果类型情况", title=True, title_lvl=3)
            calInd(rec, t_type, 'type')
            _x, _y = printTable("成果类型", t_type)
            _fn = DataHandler.doBox.doBarH(u"任务类型分布", "数量", _x, _y)
            doc.addPic(_fn, sizeof=5)

            _print("5、任务的成果归档情况", title=True, title_lvl=3)
            calInd(rec, t_base, 'base')
            _x, _y = printTable("成果归档", t_base)
            _fn = DataHandler.doBox.doBarH(u"任务归档分布", "数量", _x, _y)
            doc.addPic(_fn, sizeof=5)

        doc.saveFile('report.docx')
        # _cmd = 'python doc2pdf.py report.docx report-%s.pdf' % time.strftime('%Y%m%d', time.localtime(time.time()))
        # os.system(_cmd)

        return True

    except Exception, e:
        print e
        return False


def printInd(recs):

    for _r in recs:
        print(u"\t%s：%d" % (_r, recs[_r]))


if __name__ == '__main__':

    if main(None):
        print("任务情况：")
        printInd(task)
        print("任务来源情况：")
        printInd(source)
        print("计划情况：")
        printInd(plan)
        print("成果类型情况：")
        printInd(t_type)
        print("成果归属情况：")
        printInd(t_base)

    print("Done!")
#
# Eof
